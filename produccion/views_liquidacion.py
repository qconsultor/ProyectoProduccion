# produccion/views_liquidacion.py
from django.shortcuts import render, get_object_or_404, redirect
from django.http import JsonResponse, HttpResponse
from django.db import transaction
from django.contrib import messages
from django.db.models import Sum, F, Q
from django.views.decorators.http import require_http_methods
from django.template.loader import render_to_string

# Para exportaciones
from io import BytesIO
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models_liquidacion import Liquidacion, LiquidacionDetalle
from .forms_liquidacion import LiquidacionForm, LiquidacionDetalleFormSet
from .models import Consignacion, ConsignacionDetalle, Cliente, Producto
from .integracion_rq import insertar_prefactura_rq

# =========================
# VISTAS DE CONSULTA
# =========================

def listar_liquidaciones(request):
    """
    Lista todas las liquidaciones con filtros opcionales.
    """
    liquidaciones = Liquidacion.objects.all().order_by('-id')
    
    # Filtros opcionales
    cliente_id = request.GET.get('cliente_id')
    fecha_desde = request.GET.get('fecha_desde')
    fecha_hasta = request.GET.get('fecha_hasta')
    
    if cliente_id:
        liquidaciones = liquidaciones.filter(cliente_id=cliente_id)
    
    if fecha_desde:
        liquidaciones = liquidaciones.filter(fecha__gte=fecha_desde)
    
    if fecha_hasta:
        liquidaciones = liquidaciones.filter(fecha__lte=fecha_hasta)
    
    # Calcular total general
    total_general = sum(liq.total for liq in liquidaciones)
    
    context = {
        'liquidaciones': liquidaciones,
        'total_general': total_general,
        'cliente_id': cliente_id,
        'fecha_desde': fecha_desde,
        'fecha_hasta': fecha_hasta,
    }
    return render(request, 'produccion/liquidacion/liquidacion_list.html', context)


def ver_liquidacion(request, pk):
    """
    Muestra el detalle de una liquidación específica.
    """
    liquidacion = get_object_or_404(Liquidacion, pk=pk)
    detalles = liquidacion.detalles.all()
    
    # Obtener información del cliente desde la API/BD con filtro idsucursal=10
    try:
        cliente = Cliente.objects.using('rq').filter(
            codigo=liquidacion.cliente_id,
            idsucursal=10
        ).first()
        cliente_nombre = cliente.nombre if cliente else liquidacion.cliente_id
    except:
        cliente_nombre = liquidacion.cliente_id
    
    # Enriquecer detalles con nombres de productos
    detalles_enriquecidos = []
    for det in detalles:
        try:
            producto = Producto.objects.using('rq').filter(
                orden=det.producto_id,
                idsucursal=10
            ).first()
            producto_nombre = f"{producto.codigo} - {producto.nombre}" if producto else str(det.producto_id)
        except:
            producto_nombre = str(det.producto_id)
        
        detalles_enriquecidos.append({
            'detalle': det,
            'producto_nombre': producto_nombre
        })
    
    context = {
        'liquidacion': liquidacion,
        'detalles_enriquecidos': detalles_enriquecidos,
        'cliente_nombre': cliente_nombre,
    }
    return render(request, 'produccion/liquidacion/liquidacion_detalle.html', context)


# =========================
# VISTAS DE CREACIÓN/EDICIÓN
# =========================

# Página crear
def crear_liquidacion(request):
    if request.method == 'POST':
        # 🔍 DEBUG: Ver todos los datos POST
        print("=" * 80)
        print("📨 DATOS POST RECIBIDOS:")
        for key, value in request.POST.items():
            if key.startswith('detalles-'):
                print(f"  {key} = {value}")
        print("=" * 80)
        
        try:
            with transaction.atomic():
                # 🔹 Extraer datos del formulario principal
                cliente_id = request.POST.get('cliente_id')
                fecha = request.POST.get('fecha')
                referencia = request.POST.get('referencia', '')
                
                if not cliente_id or not fecha:
                    messages.error(request, "Cliente y fecha son obligatorios.")
                    return redirect('liquidacion:nueva')
                
                # 🔹 Crear la liquidación
                liquidacion = Liquidacion.objects.create(
                    cliente_id=cliente_id,
                    fecha=fecha,
                    referencia=referencia,
                    total=0  # Se calculará después
                )
                
                # 🔹 Procesar los detalles
                total_liquidacion = 0
                detalles_guardados = 0
                
                # Obtener el número total de formularios
                total_forms = int(request.POST.get('detalles-TOTAL_FORMS', 0))
                
                for i in range(total_forms):
                    # Extraer datos de cada detalle
                    consignacion_detalle_id = request.POST.get(f'detalles-{i}-consignacion_detalle_id')
                    producto_id = request.POST.get(f'detalles-{i}-producto_id')
                    cantidad_consignada = request.POST.get(f'detalles-{i}-cantidad_consignada', 0)
                    cantidad_devuelta = request.POST.get(f'detalles-{i}-cantidad_devuelta', 0)
                    cantidad_vendida = request.POST.get(f'detalles-{i}-cantidad_vendida', 0)
                    precio = request.POST.get(f'detalles-{i}-precio', 0)
                    
                    # 🔍 DEBUG
                    print(f"📦 Fila {i}: producto={producto_id}, vendida={cantidad_vendida}, devuelta={cantidad_devuelta}, precio={precio}")
                    
                    # Convertir a números
                    try:
                        cantidad_consignada = float(cantidad_consignada) if cantidad_consignada else 0
                        cantidad_devuelta = float(cantidad_devuelta) if cantidad_devuelta else 0
                        cantidad_vendida = float(cantidad_vendida) if cantidad_vendida else 0
                        precio = float(precio) if precio else 0
                    except (ValueError, TypeError):
                        continue
                    
                    # Solo guardar si hay movimiento (venta o devolución)
                    if cantidad_vendida > 0 or cantidad_devuelta > 0:
                        # Calcular cantidad pendiente
                        cantidad_pendiente = cantidad_consignada - cantidad_vendida - cantidad_devuelta
                        
                        # Calcular total de la línea (solo ventas generan valor)
                        total_linea = cantidad_vendida * precio
                        
                        # Crear el detalle
                        LiquidacionDetalle.objects.create(
                            liquidacion=liquidacion,
                            consignacion_detalle_id=int(consignacion_detalle_id) if consignacion_detalle_id else None,
                            producto_id=int(producto_id) if producto_id else 0,
                            cantidad_consignada=cantidad_consignada,
                            cantidad_devuelta=cantidad_devuelta,
                            cantidad_vendida=cantidad_vendida,
                            cantidad_pendiente=cantidad_pendiente,
                            precio=precio,
                            total_linea=total_linea
                        )
                        
                        total_liquidacion += total_linea
                        detalles_guardados += 1
                
                # 🔹 Actualizar el total de la liquidación
                liquidacion.total = total_liquidacion
                liquidacion.save()
                
                if detalles_guardados == 0:
                    messages.warning(request, "No se guardaron detalles. Debes ingresar al menos una venta o devolución.")
                else:
                    messages.success(request, f"✅ Liquidación #{liquidacion.id} (Ref: {liquidacion.referencia}) guardada con éxito! Total: ${total_liquidacion:.2f} ({detalles_guardados} productos)")
                    
                    # 🔹 INTEGRACIÓN CON RQ: Insertar en prefactura
                    try:
                        success, mensaje, prefactura_id = insertar_prefactura_rq(liquidacion)
                        if success:
                            messages.success(request, f"✅ {mensaje} - Lista para facturar en RQ")
                        else:
                            messages.warning(request, f"⚠️ Liquidación guardada pero no se pudo crear prefactura: {mensaje}")
                    except Exception as e:
                        messages.warning(request, f"⚠️ Liquidación guardada pero error en integración RQ: {str(e)}")
                
                # Redirigir a nueva liquidación (el correlativo se cargará automáticamente)
                return redirect('liquidacion:nueva')
                
        except Exception as e:
            messages.error(request, f"Error al guardar la liquidación: {str(e)}")
            print(f"❌ Error en crear_liquidacion: {e}")
            import traceback
            traceback.print_exc()
    
    # GET request - mostrar formulario vacío
    form = LiquidacionForm(initial={'fecha': None})
    formset = LiquidacionDetalleFormSet(prefix='detalles', queryset=LiquidacionDetalle.objects.none())

    context = {
        'form': form,
        'formset': formset,
    }
    return render(request, 'produccion/liquidacion/liquidacion_form.html', context)

# Editar
def editar_liquidacion(request, pk):
    """
    Edita una liquidación existente.
    """
    liquidacion = get_object_or_404(Liquidacion, pk=pk)
    
    if request.method == 'POST':
        try:
            with transaction.atomic():
                # Actualizar datos principales
                liquidacion.cliente_id = request.POST.get('cliente_id')
                liquidacion.fecha = request.POST.get('fecha')
                liquidacion.referencia = request.POST.get('referencia', '')
                liquidacion.consignacion_id = request.POST.get('consignacion_id') or None
                
                # Eliminar detalles anteriores
                liquidacion.detalles.all().delete()
                
                # Crear nuevos detalles desde la tabla
                total_liquidacion = 0
                detalles_data = []
                
                # Extraer datos de los productos
                idx = 0
                while True:
                    consignacion_detalle_id = request.POST.get(f'detalle_{idx}_consignacion_detalle_id')
                    if not consignacion_detalle_id:
                        break
                    
                    producto_id = request.POST.get(f'detalle_{idx}_producto_id')
                    cantidad_consignada = float(request.POST.get(f'detalle_{idx}_cantidad_consignada', 0))
                    cantidad_devuelta = float(request.POST.get(f'detalle_{idx}_devolver', 0))
                    cantidad_vendida = float(request.POST.get(f'detalle_{idx}_vender', 0))
                    precio = float(request.POST.get(f'detalle_{idx}_precio', 0))
                    
                    total_linea = cantidad_vendida * precio
                    cantidad_pendiente = cantidad_consignada - cantidad_devuelta - cantidad_vendida
                    
                    LiquidacionDetalle.objects.create(
                        liquidacion=liquidacion,
                        consignacion_detalle_id=int(consignacion_detalle_id),
                        producto_id=producto_id,
                        cantidad_consignada=cantidad_consignada,
                        cantidad_devuelta=cantidad_devuelta,
                        cantidad_vendida=cantidad_vendida,
                        cantidad_pendiente=cantidad_pendiente,
                        precio=precio,
                        total_linea=total_linea
                    )
                    
                    total_liquidacion += total_linea
                    idx += 1
                
                liquidacion.total = total_liquidacion
                liquidacion.save()
                
                messages.success(request, f'✅ Liquidación #{liquidacion.id} (Ref: {liquidacion.referencia}) actualizada con éxito!')
                return redirect('liquidacion:listar')
                
        except Exception as e:
            messages.error(request, f'❌ Error al actualizar liquidación: {str(e)}')
            print(f"❌ Error en editar_liquidacion: {e}")
            import traceback
            traceback.print_exc()
    
    # GET request - cargar datos existentes
    # Obtener detalles con información de productos
    detalles = liquidacion.detalles.all()
    detalles_json = []
    
    for det in detalles:
        try:
            producto = Producto.objects.using('rq').filter(
                orden=det.producto_id,
                idsucursal=10
            ).first()
            producto_nombre = f"{producto.codigo} - {producto.nombre}" if producto else str(det.producto_id)
        except:
            producto_nombre = str(det.producto_id)
        
        detalles_json.append({
            'consignacion_detalle_id': det.consignacion_detalle_id,
            'producto_id': det.producto_id,
            'producto_nombre': producto_nombre,
            'cantidad_consignada': float(det.cantidad_consignada),
            'cantidad_devuelta': float(det.cantidad_devuelta),
            'cantidad_vendida': float(det.cantidad_vendida),
            'cantidad_pendiente': float(det.cantidad_pendiente),
            'precio': float(det.precio),
            'total_linea': float(det.total_linea)
        })
    
    import json
    
    # Crear formulario con datos de la liquidación
    form = LiquidacionForm(initial={
        'cliente_id': liquidacion.cliente_id,
        'fecha': liquidacion.fecha,
        'referencia': liquidacion.referencia
    })
    
    context = {
        'form': form,
        'liquidacion': liquidacion,
        'detalles_json': json.dumps(detalles_json),
        'es_edicion': True,
    }
    return render(request, 'produccion/liquidacion/liquidacion_form.html', context)


# =========================
# APIs AJAX
# =========================

@require_http_methods(["GET"])
def api_siguiente_referencia(request):
    """
    Obtiene el siguiente número de referencia para liquidaciones.
    Busca el último número usado y lo incrementa en 1.
    """
    try:
        # Obtener la última liquidación
        ultima_liquidacion = Liquidacion.objects.order_by('-id').first()
        
        if ultima_liquidacion and ultima_liquidacion.referencia:
            try:
                # Intentar convertir la referencia a número y sumar 1
                ultimo_numero = int(ultima_liquidacion.referencia)
                siguiente_numero = ultimo_numero + 1
            except (ValueError, TypeError):
                # Si la referencia no es numérica, empezar desde 1
                siguiente_numero = 1
        else:
            # Si no hay liquidaciones previas, empezar desde 1
            siguiente_numero = 1
        
        return JsonResponse({
            'success': True,
            'siguiente_referencia': siguiente_numero
        })
    
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)


# =========================
# APIs AJAX (búsqueda)
# =========================
# Buscar productos (hacé que devuelva el mismo formato que usás en consignaciones)
def buscar_productos_liquidacion_api(request):
    term = request.GET.get('term', '').upper()
    results = []
    # Aquí deberías llamar a tu BD RQ o servicio que devuelve productos.
    # Ejemplo mock (cambiá por tu consulta real):
    if term:
        # simulación: devolver 10 resultados
        for i in range(1, 8):
            codigo = f"PR-{i}"
            nombre = f"PRODUCTO {i}"
            results.append({'id': codigo, 'text': f"{codigo} | {nombre}", 'precio': 1.0 * i, 'existencia': 100 - i})
    return JsonResponse({'results': results})

def get_producto_detalle_liquidacion_api(request):
    pid = request.GET.get('id')
    # llamada real: consultar el producto en RQ (por API/linked server) y devolver precio/existencia
    # aqui devolvemos ejemplo
    if not pid:
        return JsonResponse({'error': 'id required'}, status=400)
    return JsonResponse({'precio': 1.0, 'existencia': 100})

# ===========================
# 📡 API 1: CONSIGNACIONES PENDIENTES POR CLIENTE
# ===========================
def api_consignaciones_por_cliente(request):
    """
    Retorna las consignaciones de un cliente específico (idsucursal=10)
    que tienen productos con cantidades pendientes de liquidar.
    """
    cliente_id = request.GET.get('cliente_id')
    if not cliente_id:
        return JsonResponse({'error': 'Falta cliente_id'}, status=400)

    try:
        # 🔹 Obtener consignaciones del cliente
        consignaciones = Consignacion.objects.filter(
            cliente_id=cliente_id
        ).prefetch_related('detalles').order_by('-fecha')

        resultados = []
        for consig in consignaciones:
            # 🔹 Calcular cantidades pendientes
            total_pendiente = 0
            tiene_pendientes = False
            
            for detalle in consig.detalles.all():
                # Calcular cantidad liquidada de este detalle
                cantidad_liquidada = LiquidacionDetalle.objects.filter(
                    consignacion_detalle_id=detalle.id
                ).aggregate(
                    total_vendida=Sum('cantidad_vendida'),
                    total_devuelta=Sum('cantidad_devuelta')
                )
                
                vendida = cantidad_liquidada['total_vendida'] or 0
                devuelta = cantidad_liquidada['total_devuelta'] or 0
                pendiente = float(detalle.cantidad) - float(vendida) - float(devuelta)
                
                if pendiente > 0:
                    tiene_pendientes = True
                    total_pendiente += pendiente * float(detalle.precio)
            
            # Solo incluir consignaciones con productos pendientes
            if tiene_pendientes:
                resultados.append({
                    'id': consig.id,
                    'referencia': consig.referencia,
                    'fecha': consig.fecha.strftime('%d/%m/%Y'),
                    'total_original': float(consig.total),
                    'total_pendiente': round(total_pendiente, 2)
                })

        return JsonResponse({'results': resultados})
    
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


# ===========================
# 📡 API 2: DETALLE DE PRODUCTOS PENDIENTES POR CONSIGNACIÓN
# ===========================
def api_detalle_consignacion_pendiente(request):
    """
    Retorna el detalle de productos de una consignación específica,
    mostrando solo los que tienen cantidades pendientes de liquidar.
    """
    consignacion_id = request.GET.get('consignacion_id')
    if not consignacion_id:
        return JsonResponse({'error': 'Falta consignacion_id'}, status=400)

    try:
        consignacion = Consignacion.objects.get(id=consignacion_id)
        detalles = ConsignacionDetalle.objects.filter(consignacion=consignacion)
        
        # 🔹 Pre-cargar productos desde BD 'rq'
        producto_ids = [d.producto_id for d in detalles]
        productos_dict = {
            p.orden: p for p in Producto.objects.using('rq').filter(orden__in=producto_ids)
        }
        
        resultados = []
        for detalle in detalles:
            # Calcular cantidades liquidadas
            liquidaciones = LiquidacionDetalle.objects.filter(
                consignacion_detalle_id=detalle.id
            ).aggregate(
                total_vendida=Sum('cantidad_vendida'),
                total_devuelta=Sum('cantidad_devuelta')
            )
            
            vendida = float(liquidaciones['total_vendida'] or 0)
            devuelta = float(liquidaciones['total_devuelta'] or 0)
            consignada = float(detalle.cantidad)
            pendiente = consignada - vendida - devuelta
            
            # Solo incluir productos con cantidad pendiente
            if pendiente > 0:
                producto = productos_dict.get(detalle.producto_id)
                
                resultados.append({
                    'consignacion_detalle_id': detalle.id,
                    'producto_id': detalle.producto_id,
                    'producto_codigo': producto.codigo if producto else f"ID:{detalle.producto_id}",
                    'producto_nombre': producto.nombre if producto else "Producto no encontrado",
                    'cantidad_consignada': consignada,
                    'cantidad_vendida': vendida,
                    'cantidad_devuelta': devuelta,
                    'cantidad_pendiente': pendiente,
                    'precio': float(detalle.precio),
                    'total_pendiente': round(pendiente * float(detalle.precio), 2)
                })
        
        return JsonResponse({
            'consignacion': {
                'id': consignacion.id,
                'referencia': consignacion.referencia,
                'fecha': consignacion.fecha.strftime('%d/%m/%Y')
            },
            'detalles': resultados
        })
    
    except Consignacion.DoesNotExist:
        return JsonResponse({'error': 'Consignación no encontrada'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


# =========================
# EXPORTACIONES
# =========================

def exportar_pdf(request, pk):
    """Exporta una liquidación a PDF"""
    liquidacion = get_object_or_404(Liquidacion, pk=pk)
    detalles = liquidacion.detalles.all()
    
    # Obtener nombre del cliente con filtro idsucursal=10
    try:
        cliente = Cliente.objects.using('rq').filter(
            codigo=liquidacion.cliente_id,
            idsucursal=10
        ).first()
        cliente_nombre = cliente.nombre if cliente else liquidacion.cliente_id
    except:
        cliente_nombre = liquidacion.cliente_id
    
    # Crear el PDF
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []
    styles = getSampleStyleSheet()
    
    # Título
    title = Paragraph(f"<b>LIQUIDACIÓN DE CONSIGNACIÓN #{liquidacion.id}</b>", styles['Title'])
    elements.append(title)
    elements.append(Spacer(1, 0.3*inch))
    
    # Información general
    info_data = [
        ['Cliente:', f"{cliente_nombre} ({liquidacion.cliente_id})"],
        ['Fecha:', liquidacion.fecha.strftime('%d/%m/%Y')],
        ['Referencia:', liquidacion.referencia or '-'],
        ['Consignación:', str(liquidacion.consignacion_id) if liquidacion.consignacion_id else '-'],
    ]
    info_table = Table(info_data, colWidths=[2*inch, 4*inch])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    elements.append(info_table)
    elements.append(Spacer(1, 0.3*inch))
    
    # Tabla de productos
    data = [['#', 'Producto', 'Consig.', 'Devuelto', 'Vendido', 'Precio', 'Total']]
    
    for idx, det in enumerate(detalles, 1):
        # Obtener nombre del producto con filtro idsucursal=10
        try:
            producto = Producto.objects.using('rq').filter(
                orden=det.producto_id,
                idsucursal=10
            ).first()
            producto_nombre = f"{producto.codigo} - {producto.nombre}" if producto else str(det.producto_id)
        except:
            producto_nombre = str(det.producto_id)
        
        data.append([
            str(idx),
            producto_nombre[:40],
            f"{det.cantidad_consignada:.2f}",
            f"{det.cantidad_devuelta:.2f}",
            f"{det.cantidad_vendida:.2f}",
            f"${det.precio:.2f}",
            f"${det.total_linea:.2f}"
        ])
    
    # Total
    data.append(['', '', '', '', '', 'TOTAL:', f"${liquidacion.total:.2f}"])
    
    table = Table(data, colWidths=[0.5*inch, 2.5*inch, 0.8*inch, 0.8*inch, 0.8*inch, 0.8*inch, 1*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, -1), (-1, -1), colors.beige),
        ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    elements.append(table)
    
    doc.build(elements)
    buffer.seek(0)
    
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="liquidacion_{liquidacion.id}.pdf"'
    return response


def exportar_excel(request, pk):
    """Exporta una liquidación a Excel"""
    liquidacion = get_object_or_404(Liquidacion, pk=pk)
    detalles = liquidacion.detalles.all()
    
    # Obtener nombre del cliente con filtro idsucursal=10
    try:
        cliente = Cliente.objects.using('rq').filter(
            codigo=liquidacion.cliente_id,
            idsucursal=10
        ).first()
        cliente_nombre = cliente.nombre if cliente else liquidacion.cliente_id
    except:
        cliente_nombre = liquidacion.cliente_id
    
    # Crear el workbook
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"Liquidación {liquidacion.id}"
    
    # Estilos
    header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=12)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # Título
    ws.merge_cells('A1:G1')
    ws['A1'] = f'LIQUIDACIÓN DE CONSIGNACIÓN #{liquidacion.id}'
    ws['A1'].font = Font(bold=True, size=14)
    ws['A1'].alignment = Alignment(horizontal='center')
    
    # Información general
    ws['A3'] = 'Cliente:'
    ws['B3'] = f"{cliente_nombre} ({liquidacion.cliente_id})"
    ws['A4'] = 'Fecha:'
    ws['B4'] = liquidacion.fecha.strftime('%d/%m/%Y')
    ws['A5'] = 'Referencia:'
    ws['B5'] = liquidacion.referencia or '-'
    ws['A6'] = 'Consignación:'
    ws['B6'] = liquidacion.consignacion_id or '-'
    
    # Encabezados de tabla
    headers = ['#', 'Producto', 'Consignado', 'Devuelto', 'Vendido', 'Precio', 'Total']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=8, column=col)
        cell.value = header
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center')
        cell.border = border
    
    # Datos
    row = 9
    for idx, det in enumerate(detalles, 1):
        # Obtener nombre del producto con filtro idsucursal=10
        try:
            producto = Producto.objects.using('rq').filter(
                orden=det.producto_id,
                idsucursal=10
            ).first()
            producto_nombre = f"{producto.codigo} - {producto.nombre}" if producto else str(det.producto_id)
        except:
            producto_nombre = str(det.producto_id)
        
        ws.cell(row=row, column=1).value = idx
        ws.cell(row=row, column=2).value = producto_nombre
        ws.cell(row=row, column=3).value = float(det.cantidad_consignada)
        ws.cell(row=row, column=4).value = float(det.cantidad_devuelta)
        ws.cell(row=row, column=5).value = float(det.cantidad_vendida)
        ws.cell(row=row, column=6).value = float(det.precio)
        ws.cell(row=row, column=7).value = float(det.total_linea)
        
        for col in range(1, 8):
            ws.cell(row=row, column=col).border = border
        
        row += 1
    
    # Total
    ws.cell(row=row, column=6).value = 'TOTAL:'
    ws.cell(row=row, column=6).font = Font(bold=True)
    ws.cell(row=row, column=7).value = float(liquidacion.total)
    ws.cell(row=row, column=7).font = Font(bold=True)
    
    # Ajustar anchos de columna
    ws.column_dimensions['A'].width = 5
    ws.column_dimensions['B'].width = 40
    ws.column_dimensions['C'].width = 12
    ws.column_dimensions['D'].width = 12
    ws.column_dimensions['E'].width = 12
    ws.column_dimensions['F'].width = 12
    ws.column_dimensions['G'].width = 12
    
    # Guardar en buffer
    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="liquidacion_{liquidacion.id}.xlsx"'
    return response


def exportar_word(request, pk):
    """Exporta una liquidación a Word"""
    liquidacion = get_object_or_404(Liquidacion, pk=pk)
    detalles = liquidacion.detalles.all()
    
    # Obtener nombre del cliente con filtro idsucursal=10
    try:
        cliente = Cliente.objects.using('rq').filter(
            codigo=liquidacion.cliente_id,
            idsucursal=10
        ).first()
        cliente_nombre = cliente.nombre if cliente else liquidacion.cliente_id
    except:
        cliente_nombre = liquidacion.cliente_id
    
    # Crear documento
    doc = Document()
    
    # Título
    title = doc.add_heading(f'LIQUIDACIÓN DE CONSIGNACIÓN #{liquidacion.id}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Información general
    doc.add_paragraph(f'Cliente: {cliente_nombre} ({liquidacion.cliente_id})')
    doc.add_paragraph(f'Fecha: {liquidacion.fecha.strftime("%d/%m/%Y")}')
    doc.add_paragraph(f'Referencia: {liquidacion.referencia or "-"}')
    doc.add_paragraph(f'Consignación: {liquidacion.consignacion_id or "-"}')
    doc.add_paragraph('')  # Espacio
    
    # Tabla de productos
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Light Grid Accent 1'
    
    # Encabezados
    headers = ['#', 'Producto', 'Consignado', 'Devuelto', 'Vendido', 'Precio', 'Total']
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
    
    # Datos
    for idx, det in enumerate(detalles, 1):
        # Obtener nombre del producto con filtro idsucursal=10
        try:
            producto = Producto.objects.using('rq').filter(
                orden=det.producto_id,
                idsucursal=10
            ).first()
            producto_nombre = f"{producto.codigo} - {producto.nombre}" if producto else str(det.producto_id)
        except:
            producto_nombre = str(det.producto_id)
        
        row = table.add_row()
        row.cells[0].text = str(idx)
        row.cells[1].text = producto_nombre[:50]
        row.cells[2].text = f"{det.cantidad_consignada:.2f}"
        row.cells[3].text = f"{det.cantidad_devuelta:.2f}"
        row.cells[4].text = f"{det.cantidad_vendida:.2f}"
        row.cells[5].text = f"${det.precio:.2f}"
        row.cells[6].text = f"${det.total_linea:.2f}"
    
    # Total
    row = table.add_row()
    row.cells[5].text = 'TOTAL:'
    row.cells[5].paragraphs[0].runs[0].font.bold = True
    row.cells[6].text = f"${liquidacion.total:.2f}"
    row.cells[6].paragraphs[0].runs[0].font.bold = True
    
    # Guardar en buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="liquidacion_{liquidacion.id}.docx"'
    return response
